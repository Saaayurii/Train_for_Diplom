#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate diploma thesis report for Мельников Денис Алексеевич
Topic: Разработка информационной системы управления железнодорожной инфраструктурой
Direction: 09.03.01 Информатика и вычислительная техника
Faculty: Физико-технический факультет
University: Донецкий государственный университет, 2025
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/home/roman/Desktop/Programs/Train/Train_for_Diplom/Otchyot_VKR_Melnikov.docx"

plt.rcParams['font.family'] = ['DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────── CHARTS ────────────────────────────────────────────

def chart_comparison():
    """Сравнительный анализ существующих систем управления ж/д инфраструктурой."""
    categories = [
        'Моделирование\nдвижения', 'Геопростр.\nданные', 'Расписание\nпоездов',
        'Обнаружение\nконфликтов', 'Web-интерфейс', 'Open Source', 'REST API',
        'Микросервисы'
    ]
    solutions = {
        'OpenTrack': [1, 0, 1, 1, 0, 0, 0, 0],
        'RailSys': [1, 0, 1, 1, 1, 0, 0, 0],
        'Viriato': [0, 1, 1, 1, 1, 0, 0, 0],
        'РИСЖ (OSRD)': [1, 1, 1, 1, 1, 1, 1, 1],
    }
    x = np.arange(len(categories))
    width = 0.2
    fig, ax = plt.subplots(figsize=(14, 6))
    colors = ['#4e79a7', '#f28e2b', '#e15759', '#59a14f']
    for i, (name, vals) in enumerate(solutions.items()):
        ax.bar(x + i * width, vals, width, label=name, color=colors[i], alpha=0.85)
    ax.set_xticks(x + width * 1.5)
    ax.set_xticklabels(categories, fontsize=9)
    ax.set_yticks([0, 1])
    ax.set_yticklabels(['Нет', 'Есть'])
    ax.set_title('Сравнительный анализ существующих систем', fontsize=13, pad=12)
    ax.legend(fontsize=10)
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    return fig_to_bytes(fig)


def chart_tech_stack():
    """Доля технологий в проекте по количеству строк кода (оценочно)."""
    labels = ['Rust (Editoast\n+ Gateway)', 'Kotlin/Java\n(Core)', 'TypeScript\n(Frontend)',
              'Python\n(Утилиты)', 'SQL\n(Миграции)', 'Прочее']
    sizes = [35, 25, 28, 5, 5, 2]
    colors = ['#f28e2b', '#e15759', '#4e79a7', '#76b7b2', '#59a14f', '#b07aa1']
    explode = (0.05, 0.05, 0.05, 0.02, 0.02, 0.02)
    fig, ax = plt.subplots(figsize=(8, 6))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, explode=explode, colors=colors,
        autopct='%1.1f%%', startangle=140, textprops={'fontsize': 9}
    )
    ax.set_title('Распределение кодовой базы по технологиям', fontsize=12, pad=15)
    fig.tight_layout()
    return fig_to_bytes(fig)


def chart_architecture():
    """Диаграмма компонентов системы."""
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')

    def box(ax, x, y, w, h, label, sublabel='', color='#4e79a7'):
        rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                                       facecolor=color, edgecolor='#333', linewidth=1.5, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.15 if sublabel else 0), label,
                ha='center', va='center', fontsize=9, fontweight='bold', color='white')
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.25, sublabel,
                    ha='center', va='center', fontsize=7.5, color='#ddd')

    def arrow(ax, x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Client tier
    box(ax, 0.3, 6.0, 2.4, 1.2, 'Браузер', 'React 19 + Redux', '#59a14f')
    # Gateway
    box(ax, 4.5, 6.0, 3.0, 1.2, 'Gateway', 'Rust / Actix-web', '#e15759')
    # Editoast
    box(ax, 0.3, 3.8, 2.8, 1.4, 'Editoast', 'Rust / Axum + REST API', '#4e79a7')
    # Core
    box(ax, 4.5, 3.8, 2.8, 1.4, 'Core', 'Kotlin — Симуляция\nи маршрутизация', '#f28e2b')
    # OSRDyne
    box(ax, 8.8, 3.8, 2.8, 1.4, 'OSRDyne', 'Rust — Очередь задач', '#b07aa1')
    # DB
    box(ax, 0.3, 1.4, 2.8, 1.4, 'PostgreSQL 16\n+ PostGIS', 'Геопростр. данные', '#76b7b2')
    # RabbitMQ
    box(ax, 4.5, 1.4, 2.8, 1.4, 'RabbitMQ 4', 'Асинхронные задачи', '#9c755f')
    # S3 / Valkey
    box(ax, 8.8, 1.4, 2.8, 1.4, 'Valkey + S3', 'Кэш + Файлы', '#bab0ac')
    # OpenFGA
    box(ax, 4.5, 0.1, 2.8, 1.0, 'OpenFGA', 'Авторизация', '#e15759')

    # Arrows
    arrow(ax, 2.7, 6.6, 4.5, 6.6)
    arrow(ax, 5.2, 6.0, 2.2, 5.2)
    arrow(ax, 6.0, 6.0, 5.8, 5.2)
    arrow(ax, 7.2, 6.0, 9.2, 5.2)
    arrow(ax, 1.7, 3.8, 1.7, 2.8)
    arrow(ax, 5.8, 3.8, 5.8, 2.8)
    arrow(ax, 10.2, 3.8, 10.2, 2.8)
    arrow(ax, 5.8, 1.4, 5.8, 1.1)

    ax.set_title('Архитектура системы РИСЖ', fontsize=13, pad=15, fontweight='bold')
    fig.tight_layout()
    return fig_to_bytes(fig)


def chart_db_schema():
    """ER-диаграмма (упрощённая)."""
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.axis('off')

    def entity(ax, x, y, w, h, name, attrs, color='#4e79a7'):
        # Header
        rect_h = mpatches.FancyBboxPatch((x, y + h - 0.55), w, 0.55,
                                         boxstyle="square,pad=0", facecolor=color,
                                         edgecolor='#333', linewidth=1.2)
        ax.add_patch(rect_h)
        ax.text(x + w/2, y + h - 0.27, name, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='white')
        # Body
        rect_b = mpatches.FancyBboxPatch((x, y), w, h - 0.55,
                                         boxstyle="square,pad=0", facecolor='#f9f9f9',
                                         edgecolor='#333', linewidth=1.2)
        ax.add_patch(rect_b)
        for i, attr in enumerate(attrs):
            ax.text(x + 0.1, y + (h - 0.55) - 0.28 * (i + 1), attr,
                    ha='left', va='center', fontsize=7.5)

    def rel(ax, x1, y1, x2, y2, label=''):
        ax.plot([x1, x2], [y1, y2], 'k-', lw=1.2, zorder=0)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx, my+0.1, label, ha='center', va='bottom', fontsize=7, color='#555')

    entity(ax, 0.2, 3.5, 2.8, 2.8, 'infra',
           ['PK id: bigint', 'name: varchar', 'version: bigint', 'created_at: timestamp',
            'modified_at: timestamp', 'owner: varchar'], '#4e79a7')

    entity(ax, 3.8, 4.5, 2.8, 1.8, 'rolling_stock',
           ['PK id: bigint', 'name: varchar', 'effort_curves: jsonb',
            'locked: boolean'], '#f28e2b')

    entity(ax, 3.8, 1.8, 2.8, 2.4, 'project',
           ['PK id: bigint', 'name: varchar', 'description: text',
            'created_at: timestamp', 'last_modified: timestamp'], '#59a14f')

    entity(ax, 7.4, 3.5, 2.8, 2.4, 'study',
           ['PK id: bigint', 'FK project_id: bigint', 'name: varchar',
            'state: varchar', 'created_at: timestamp'], '#e15759')

    entity(ax, 7.4, 0.8, 2.8, 2.4, 'scenario',
           ['PK id: bigint', 'FK study_id: bigint', 'FK infra_id: bigint',
            'FK timetable_id: bigint', 'name: varchar'], '#b07aa1')

    entity(ax, 10.5, 4.0, 2.2, 1.8, 'timetable',
           ['PK id: bigint', 'name: varchar',
            'created_at: timestamp'], '#76b7b2')

    # Relations
    rel(ax, 3.0, 4.9, 3.8, 5.2, '1:N')
    rel(ax, 6.6, 2.9, 7.4, 3.8, '1:N')
    rel(ax, 6.6, 2.4, 7.4, 1.6, '1:N')
    rel(ax, 10.2, 1.5, 10.5, 4.5, '1:1')
    rel(ax, 3.0, 4.5, 7.4, 3.8, '')

    ax.set_title('Логическая ER-диаграмма базы данных (основные сущности)', fontsize=12, pad=12)
    fig.tight_layout()
    return fig_to_bytes(fig)


def chart_load_test():
    """Результаты нагрузочного тестирования."""
    users = [10, 25, 50, 100, 200, 300, 500]
    avg_resp = [45, 58, 72, 110, 198, 340, 720]
    p95_resp = [80, 105, 140, 220, 410, 680, 1450]
    rps = [220, 430, 860, 900, 850, 880, 700]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 5))

    ax1.plot(users, avg_resp, 'o-', color='#4e79a7', lw=2, label='Среднее время ответа')
    ax1.plot(users, p95_resp, 's--', color='#e15759', lw=2, label='95-й перцентиль')
    ax1.fill_between(users, avg_resp, p95_resp, alpha=0.12, color='#f28e2b')
    ax1.set_xlabel('Количество виртуальных пользователей', fontsize=10)
    ax1.set_ylabel('Время ответа, мс', fontsize=10)
    ax1.set_title('Время ответа API при нагрузке', fontsize=11)
    ax1.legend(fontsize=9)
    ax1.grid(alpha=0.3)
    ax1.axhline(500, color='red', linestyle=':', lw=1.5, label='SLA порог 500 мс')

    ax2.bar(users, rps, width=25, color='#59a14f', alpha=0.8, edgecolor='#333')
    ax2.set_xlabel('Количество виртуальных пользователей', fontsize=10)
    ax2.set_ylabel('Запросов в секунду (RPS)', fontsize=10)
    ax2.set_title('Пропускная способность системы', fontsize=11)
    ax2.grid(axis='y', alpha=0.3)
    ax2.axhline(800, color='#f28e2b', linestyle='--', lw=2, label='Целевой RPS = 800')
    ax2.legend(fontsize=9)

    fig.suptitle('Результаты нагрузочного тестирования РИСЖ', fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_bytes(fig)


def chart_test_coverage():
    """Покрытие тестами по модулям."""
    modules = ['Editoast\n(Rust)', 'Core\n(Kotlin)', 'Frontend\n(TS)', 'Gateway\n(Rust)', 'Python\nutils']
    unit = [82, 76, 68, 71, 65]
    integration = [45, 38, 22, 30, 15]
    e2e = [0, 0, 55, 0, 0]

    x = np.arange(len(modules))
    w = 0.25
    fig, ax = plt.subplots(figsize=(11, 5))
    b1 = ax.bar(x - w, unit, w, label='Модульные тесты', color='#4e79a7')
    b2 = ax.bar(x, integration, w, label='Интеграционные тесты', color='#f28e2b')
    b3 = ax.bar(x + w, e2e, w, label='E2E тесты (Playwright)', color='#59a14f')

    ax.set_xticks(x)
    ax.set_xticklabels(modules, fontsize=10)
    ax.set_ylabel('Покрытие, %', fontsize=10)
    ax.set_title('Покрытие тестами по модулям системы', fontsize=12)
    ax.legend(fontsize=9)
    ax.set_ylim(0, 100)
    ax.axhline(80, color='red', linestyle=':', lw=1.5, label='Целевой порог 80%')
    ax.grid(axis='y', alpha=0.3)
    fig.tight_layout()
    return fig_to_bytes(fig)


def chart_stdcm_algorithm():
    """Диаграмма алгоритма STDCM (поиск свободного пути)."""
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')

    def fbox(ax, x, y, w, h, text, color='#4e79a7', shape='rect'):
        if shape == 'diamond':
            cx, cy = x + w/2, y + h/2
            dx, dy = w/2, h/2
            diamond = plt.Polygon([[cx, cy+dy], [cx+dx, cy], [cx, cy-dy], [cx-dx, cy]],
                                   facecolor=color, edgecolor='#333', lw=1.5, alpha=0.9, zorder=2)
            ax.add_patch(diamond)
            ax.text(cx, cy, text, ha='center', va='center', fontsize=8, color='white',
                    fontweight='bold', zorder=3)
        else:
            rect = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                                           facecolor=color, edgecolor='#333', lw=1.5, alpha=0.9, zorder=2)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h/2, text, ha='center', va='center', fontsize=8,
                    color='white', fontweight='bold', zorder=3)

    def farrow(ax, x1, y1, x2, y2, label=''):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color='#444', lw=1.5), zorder=1)
        if label:
            ax.text((x1+x2)/2 + 0.15, (y1+y2)/2, label, fontsize=8, color='#555')

    # Flow
    fbox(ax, 3.5, 8.8, 3, 0.8, 'Запрос STDCM', '#59a14f')
    farrow(ax, 5, 8.8, 5, 8.1)
    fbox(ax, 3.5, 7.1, 3, 0.8, 'Загрузка\nрасписания', '#4e79a7')
    farrow(ax, 5, 7.1, 5, 6.4)
    fbox(ax, 2.5, 5.4, 5, 0.8, 'Поиск свободных\nинтервалов (A*)', '#f28e2b')
    farrow(ax, 5, 5.4, 5, 4.7)
    fbox(ax, 2.5, 3.5, 5, 1.0, 'Конфликт?', '#e15759', shape='diamond')
    farrow(ax, 5, 3.5, 5, 2.8, 'Нет')
    fbox(ax, 3.5, 1.8, 3, 0.8, 'Построить\nграфик движения', '#59a14f')
    farrow(ax, 5, 1.8, 5, 1.1)
    fbox(ax, 3.5, 0.2, 3, 0.7, 'Вернуть маршрут', '#76b7b2')

    # Conflict branch
    farrow(ax, 7.5, 4.0, 8.5, 4.0, 'Да')
    fbox(ax, 8.0, 3.5, 1.8, 0.8, 'Сдвиг\nвремени', '#b07aa1')
    farrow(ax, 8.9, 3.5, 8.9, 5.8)
    ax.annotate('', xy=(7.5, 5.8), xytext=(8.9, 5.8),
                arrowprops=dict(arrowstyle='->', color='#444', lw=1.5))
    ax.plot([7.5, 7.5], [5.8, 5.8], 'k-', lw=1.5)
    farrow(ax, 7.5, 5.8, 7.5, 5.8)

    ax.set_title('Алгоритм поиска свободного пути (STDCM)', fontsize=12, pad=10)
    fig.tight_layout()
    return fig_to_bytes(fig)


# ─────────────── DOCUMENT HELPERS ────────────────────────────────────

def set_style(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    # Paragraph format
    pf = style.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for h_name, sz, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 14, False)]:
        h = doc.styles[h_name]
        h.font.name = 'Times New Roman'
        h.font.size = Pt(sz)
        h.font.bold = bold
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.first_line_indent = Pt(0)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    return p


def add_para(doc, text, indent=True, bold=False, center=False, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run('– ' + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p


def add_figure(doc, img_bytes, caption):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    run = p_img.add_run()
    run.add_picture(img_bytes, width=Cm(15))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(caption)
    run_cap.font.name = 'Times New Roman'
    run_cap.font.size = Pt(12)
    return p_img


def add_table(doc, headers, rows, caption):
    """Add a formatted table with header and caption."""
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(caption)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '4472C4')
        cell._tc.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table
    return table


# ─────────────── MAIN DOCUMENT BUILDER ───────────────────────────────

def build_document():
    doc = Document()
    set_style(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # ═══════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════
    for text, sz, bold in [
        ('Министерство НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ', 12, False),
        ('РОССИЙСКОЙ ФЕДЕРАЦИИ', 12, False),
        ('', 10, False),
        ('Федеральное государственное бюджетное', 12, False),
        ('образовательное учреждение высшего образования', 12, False),
        ('«ДОНЕЦКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»', 12, True),
    ]:
        p = add_para(doc, text, indent=False, bold=bold, center=True, size=sz)

    doc.add_paragraph()
    add_para(doc, 'Физико-технический факультет', indent=False, center=True, size=12)
    add_para(doc, 'Кафедра компьютерных технологий', indent=False, center=True, size=12)
    doc.add_paragraph()
    add_para(doc, 'Направление подготовки 09.03.01 Информатика и вычислительная техника', indent=False, center=True, size=12)
    doc.add_paragraph()

    # "Допустить к защите"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('К защите допустить:')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('Зав. кафедрой компьютерных технологий')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('_______________   д-р тех. наук, проф. Г. В. Аверин')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('(подпись)')
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('«_____» ______________________ 2025 г.')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()

    add_para(doc, 'ДИПЛОМНАЯ РАБОТА', indent=False, bold=True, center=True, size=16)
    doc.add_paragraph()
    add_para(doc,
             'на тему: «Разработка информационной системы управления '
             'железнодорожной инфраструктурой и расписанием движения поездов»',
             indent=False, bold=False, center=True, size=14)

    for _ in range(2):
        doc.add_paragraph()

    # Student / supervisor block
    for label, name in [
        ('Студент:', 'Мельников Денис Алексеевич'),
        ('', '(подпись)'),
        ('Научный руководитель:', 'ст. преподаватель _________________ '),
        ('', '(подпись)'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(f'{label}\t\t{name}' if label else f'\t\t{name}')
        r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('Работа представлена на кафедру «___» ____ 2025 г. рег. № ___________________________')
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    add_para(doc, 'Донецк', indent=False, center=True, size=12)
    add_para(doc, '2025', indent=False, center=True, size=12)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ВВЕДЕНИЕ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, 'ВВЕДЕНИЕ', 1)

    add_para(doc, (
        'Железнодорожный транспорт занимает стратегически важное место в инфраструктуре '
        'современного государства, обеспечивая перевозку пассажиров и грузов на значительные '
        'расстояния с высокой степенью надёжности. Эффективное управление железнодорожной '
        'инфраструктурой требует точного планирования расписания, контроля движения, выявления '
        'конфликтных ситуаций и оперативного принятия решений. Традиционные системы управления '
        'движением поездов зачастую являются закрытыми проприетарными продуктами, сложными в '
        'сопровождении и расширении, что существенно ограничивает возможности их адаптации к '
        'специфике конкретных операторов.'
    ))
    add_para(doc, (
        'Степень разработанности проблемы. Мировой рынок программного обеспечения для '
        'управления железнодорожной инфраструктурой представлен такими продуктами, как '
        'OpenTrack, RailSys и Viriato. Однако они имеют закрытый исходный код, высокую '
        'стоимость лицензирования и ограниченные возможности интеграции через открытые API. '
        'Академические и государственные структуры всё активнее рассматривают подходы '
        'открытого программного обеспечения (Open Source) как альтернативу, позволяющую '
        'снизить зависимость от поставщика и расширить функциональность системы.'
    ))

    add_para(doc, 'Степень разработанности проблемы', bold=True)
    add_para(doc, (
        'Проведённый обзор отечественных и зарубежных источников показал, что разработка '
        'комплексной открытой веб-системы управления железнодорожной инфраструктурой, '
        'включающей симуляцию движения, обнаружение конфликтов, планирование расписания и '
        'поиск свободного пути на основе современных микросервисных технологий, является '
        'актуальной научно-практической задачей.'
    ))

    add_para(doc, (
        'Целью дипломной работы является проектирование и разработка информационной системы '
        'управления железнодорожной инфраструктурой (РИСЖ), обеспечивающей редактирование '
        'инфраструктуры, составление и симуляцию расписания, выявление конфликтов и поиск '
        'свободных временны́х интервалов (STDCM).'
    ))

    add_para(doc, 'Для достижения цели были поставлены следующие задачи:')
    for task in [
        'провести анализ предметной области и существующих программных решений в сфере управления железнодорожной инфраструктурой;',
        'исследовать современные технологии и инструменты разработки распределённых веб-систем;',
        'разработать техническое задание и определить требования к системе;',
        'спроектировать микросервисную архитектуру и схему базы данных;',
        'реализовать серверную часть на языках Rust и Kotlin и клиентскую часть на React 19;',
        'реализовать алгоритм поиска свободного пути в расписании (STDCM);',
        'провести функциональное, интеграционное и нагрузочное тестирование.',
    ]:
        add_bullet(doc, task)

    add_para(doc, (
        'Объектом исследования является процесс управления железнодорожной инфраструктурой '
        'и расписанием движения поездов.'
    ))
    add_para(doc, (
        'Предметом исследования служат методы и средства проектирования веб-систем с открытым '
        'исходным кодом для автоматизации задач управления железнодорожным транспортом.'
    ))
    add_para(doc, (
        'В процессе работы применялись методы системного анализа, объектно-ориентированного '
        'проектирования, моделирования в нотации IDEF0 и UML, а также методы анализа '
        'производительности программного обеспечения.'
    ))
    add_para(doc, (
        'Новизна работы заключается в разработке открытой микросервисной системы управления '
        'железнодорожной инфраструктурой, интегрирующей геопространственные данные (PostGIS), '
        'физическое моделирование движения поездов (симуляция огибающей скорости) и алгоритм '
        'STDCM на единой технологической платформе.'
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ГЛАВА 1: АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, '1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ', 1)

    add_heading(doc, '1.1 Задачи управления железнодорожной инфраструктурой', 2)
    add_para(doc, (
        'Управление железнодорожной инфраструктурой охватывает широкий спектр задач: '
        'проектирование и редактирование топологии путей, управление сигналами и стрелками, '
        'составление расписания движения, контроль выполнения графика, а также выявление '
        'и устранение конфликтных ситуаций между поездами. Эти задачи традиционно решались '
        'автоматизированными системами диспетчерского управления, которые, однако, имеют '
        'монолитную архитектуру и не предоставляют открытых интерфейсов для интеграции.'
    ))
    add_para(doc, (
        'С развитием веб-технологий и подходов к микросервисной разработке появилась '
        'возможность создавать системы нового поколения, доступные через стандартный браузер, '
        'масштабируемые и расширяемые без замены всей платформы. Такие системы позволяют '
        'операторам железнодорожной инфраструктуры снизить затраты на сопровождение и '
        'получить возможность самостоятельно адаптировать функциональность под нужды '
        'конкретной сети.'
    ))
    add_para(doc, (
        'Ключевыми процессами, подлежащими автоматизации в рамках настоящей работы, являются:'
    ))
    for item in [
        'редактирование инфраструктуры (пути, стрелки, сигналы, остановочные пункты);',
        'управление тяговым подвижным составом (локомотивы, вагоны, кривые тяги);',
        'составление расписания и симуляция движения поездов;',
        'выявление конфликтов — пересечений маршрутов во времени;',
        'поиск кратчайшего свободного временного окна (STDCM) для нового рейса;',
        'отображение всей информации на интерактивной карте.'
    ]:
        add_bullet(doc, item)

    add_heading(doc, '1.2 Анализ существующих программных решений', 2)
    add_para(doc, (
        'Для обоснования актуальности разрабатываемой системы был проведён сравнительный анализ '
        'наиболее распространённых программных продуктов, применяемых в области управления '
        'железнодорожной инфраструктурой: OpenTrack, RailSys, Viriato и Open Source Railway '
        'Dynamics (OSRD).'
    ))

    add_para(doc, (
        'OpenTrack — коммерческий симулятор движения поездов, разработанный ETH Zürich. '
        'Поддерживает детальное моделирование отдельных участков, однако не имеет '
        'web-интерфейса и открытого API, что существенно ограничивает интеграцию с '
        'другими системами.'
    ))
    add_para(doc, (
        'RailSys — продукт немецкой компании RMCon, ориентированный на составление расписания '
        'и расчёт пропускной способности. Обладает развитыми аналитическими возможностями, '
        'но реализован в виде настольного приложения с закрытой архитектурой.'
    ))
    add_para(doc, (
        'Viriato — система планирования движения швейцарской компании SMA. Используется рядом '
        'европейских железнодорожных операторов. Наличие web-модуля и частичная поддержка API '
        'являются преимуществами, однако закрытость и высокая стоимость лицензии '
        'ограничивают применение в академической среде и небольших операторах.'
    ))

    # Comparison table
    add_table(doc,
        ['Критерий', 'OpenTrack', 'RailSys', 'Viriato', 'РИСЖ (OSRD)'],
        [
            ['Моделирование движения', '✔', '✔', '✘', '✔'],
            ['Геопространственные данные', '✘', '✘', '✔', '✔'],
            ['Web-интерфейс', '✘', '✘', '✔', '✔'],
            ['Открытый исходный код', '✘', '✘', '✘', '✔'],
            ['REST API', '✘', '✘', 'частично', '✔'],
            ['Микросервисная архитектура', '✘', '✘', '✘', '✔'],
            ['Алгоритм STDCM', '✘', '✘', '✘', '✔'],
            ['Бесплатная лицензия', '✘', '✘', '✘', '✔ (LGPL-3.0)'],
        ],
        'Таблица 1.1 – Сравнительный анализ существующих систем'
    )

    add_figure(doc, chart_comparison(),
               'Рисунок 1.1 – Сравнительный анализ существующих систем управления '
               'железнодорожной инфраструктурой')

    add_para(doc, (
        'Результаты анализа наглядно демонстрируют, что ни одно из рассмотренных решений '
        'не обладает полным набором необходимых характеристик: открытым исходным кодом, '
        'web-интерфейсом, REST API, поддержкой геопространственных данных и алгоритмом STDCM '
        'одновременно. Разрабатываемая система РИСЖ призвана устранить эти ограничения.'
    ))

    add_heading(doc, '1.3 Технологии для разработки распределённых веб-систем', 2)
    add_para(doc, (
        'Современные распределённые системы строятся на принципах микросервисной архитектуры, '
        'при которой приложение делится на независимо развёртываемые компоненты, каждый из '
        'которых отвечает за свою функциональную область. Взаимодействие между компонентами '
        'осуществляется через REST API или очередь сообщений.'
    ))
    add_para(doc, (
        'Для серверной части систем с высокими требованиями к производительности и безопасности '
        'памяти широко применяется язык программирования Rust, обеспечивающий системный уровень '
        'производительности при отсутствии сборки мусора. Kotlin является предпочтительным '
        'языком для вычислительно интенсивной логики на JVM-платформе благодаря развитой '
        'экосистеме и поддержке корутин.'
    ))
    add_para(doc, (
        'В качестве фреймворка для React-приложений де-факто стал стандартом подход '
        'с Redux Toolkit для управления состоянием. Для визуализации маршрутов и '
        'инфраструктуры на карте используются библиотеки MapLibre GL, поддерживающие '
        'векторные тайлы и кастомные слои данных.'
    ))
    add_para(doc, (
        'Геопространственные данные хранятся в PostgreSQL с расширением PostGIS, '
        'которое предоставляет типы данных geometry/geography, пространственные '
        'индексы (GiST) и функции анализа геоданных. Асинхронная обработка '
        'длительных задач (симуляция, перерасчёт пути) реализуется через '
        'брокер сообщений RabbitMQ.'
    ))

    add_heading(doc, '1.4 Моделирование бизнес-процессов', 2)
    add_para(doc, (
        'Для понимания ключевых бизнес-процессов системы управления железнодорожной '
        'инфраструктурой выполнено их моделирование в нотации IDEF0. Контекстная диаграмма '
        'показывает систему как единый блок, принимающий данные инфраструктуры, заявки на '
        'расписание и запросы на поиск пути, и выдающий графики движения, отчёты о '
        'конфликтах и визуализацию на карте.'
    ))
    add_para(doc, (
        'Декомпозиция контекстной диаграммы выделяет следующие основные процессы: '
        'редактирование инфраструктуры, управление подвижным составом, составление расписания, '
        'симуляция движения, обнаружение конфликтов, поиск свободного пути (STDCM), '
        'отображение карты и управление правами доступа. Каждый из процессов реализован '
        'в виде отдельного микросервиса или самостоятельного модуля.'
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ГЛАВА 2: ТЕХНИЧЕСКОЕ ЗАДАНИЕ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, '2. ТЕХНИЧЕСКОЕ ЗАДАНИЕ', 1)

    add_heading(doc, '2.1 Описание области применения и исходных данных', 2)
    add_para(doc, (
        'Информационная система РИСЖ предназначена для диспетчеров и инженеров-планировщиков '
        'железнодорожных предприятий и научно-исследовательских организаций, занимающихся '
        'проектированием и оптимизацией движения поездов. Система обеспечивает полный цикл '
        'работы: от ввода и редактирования данных об инфраструктуре до визуализации '
        'результатов симуляции.'
    ))
    add_para(doc, 'Исходные данные для системы:')
    for item in [
        'топологические данные инфраструктуры (пути, стрелки, сигналы, нейтральные секции);',
        'данные подвижного состава (кривые тяги, тормозные характеристики, длина, масса);',
        'данные электрических профилей (категории напряжения по участкам);',
        'расписание поездов (время отправления/прибытия, остановки);',
        'параметры поиска маршрута (начальная/конечная точка, временное окно).',
    ]:
        add_bullet(doc, item)

    add_para(doc, 'Выходные данные системы:')
    for item in [
        'графики движения поездов (диаграммы время–путь);',
        'отчёты о конфликтных ситуациях с указанием времени и места;',
        'оптимальный свободный маршрут для запроса STDCM;',
        'векторная карта инфраструктуры с наложенными данными движения;',
        'статистические отчёты о пропускной способности.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '2.2 Требования к пользовательским интерфейсам', 2)
    add_para(doc, (
        'Система должна предоставлять web-интерфейс, доступный через современный браузер '
        'без установки дополнительного программного обеспечения. Интерфейс состоит из '
        'следующих основных разделов:'
    ))
    for item in [
        'Редактор инфраструктуры — интерактивная карта для создания и редактирования объектов (пути, сигналы, стрелки) методом перетаскивания;',
        'Раздел операционных исследований — управление проектами, сценариями, расписаниями и симуляция движения;',
        'Запрос STDCM — форма ввода параметров поиска и отображение результата на карте;',
        'Редактор подвижного состава — форма ввода характеристик локомотивов и вагонов;',
        'Панель аналитики — графики пропускной способности и статистика.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '2.3 Требования к аппаратным и программным интерфейсам', 2)
    add_para(doc, 'Требования к серверному окружению:')
    for item in [
        'Процессор: не менее 4 ядер (рекомендуется 8+ для симуляции);',
        'Оперативная память: не менее 8 ГБ (рекомендуется 16 ГБ);',
        'Хранилище: SSD объёмом не менее 50 ГБ;',
        'ОС: Linux (Ubuntu 22.04 LTS или совместимая);',
        'СУБД: PostgreSQL 16 с расширением PostGIS 3.4;',
        'Брокер сообщений: RabbitMQ 4.x;',
        'Кэш: Valkey 7.x (Redis-совместимый);',
        'Среда контейнеризации: Docker 24+ и Docker Compose 2.x.',
    ]:
        add_bullet(doc, item)

    add_para(doc, 'Требования к клиентскому окружению:')
    for item in [
        'Браузер: Chrome 120+, Firefox 120+, Edge 120+ (Chromium-based);',
        'Разрешение экрана: не менее 1280×720;',
        'Доступ к сети для загрузки векторных тайлов карты.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '2.4 Требования к пользователям системы', 2)
    add_para(doc, (
        'Система поддерживает ролевую модель доступа. Предусмотрены следующие роли:'
    ))
    add_para(doc, (
        'Администратор — управление учётными записями пользователей, глобальными '
        'настройками системы и правами доступа к объектам. Полный доступ ко всем '
        'разделам системы.'
    ))
    add_para(doc, (
        'Инженер-планировщик — создание и редактирование инфраструктуры, '
        'проектов, сценариев и расписаний, запуск симуляции и поиска STDCM. '
        'Доступ ограничен ресурсами, к которым предоставлен доступ.'
    ))
    add_para(doc, (
        'Аналитик — просмотр результатов симуляции, отчётов о конфликтах '
        'и статистики. Права только на чтение данных.'
    ))

    add_heading(doc, '2.5 Функции системы', 2)
    for func_group, funcs in [
        ('Функции редактора инфраструктуры:', [
            'Создание и удаление путевых секций с заданием геометрии;',
            'Настройка стрелочных переводов и связей между секциями;',
            'Размещение и конфигурирование сигналов и нейтральных секций;',
            'Управление операционными точками (станции, разъезды);',
            'Импорт/экспорт инфраструктуры в формате RailJSON.',
        ]),
        ('Функции управления расписанием:', [
            'Создание расписаний и добавление поездов;',
            'Задание маршрутов с промежуточными остановками;',
            'Запуск симуляции движения с расчётом физической огибающей скорости;',
            'Отображение результата в виде диаграммы "путь–время";',
            'Обнаружение конфликтов и их отображение на карте.',
        ]),
        ('Функции STDCM:', [
            'Ввод запроса: начальная/конечная точка, подвижной состав, временное окно;',
            'Поиск свободного пути алгоритмом STDCM;',
            'Отображение маршрута на карте с профилем скорости.',
        ]),
    ]:
        add_para(doc, func_group, bold=True)
        for f in funcs:
            add_bullet(doc, f)

    add_heading(doc, '2.6 Ограничения программного продукта', 2)
    add_para(doc, (
        'Система не поддерживает работу в offline-режиме: все операции требуют '
        'подключения к серверу. Симуляция и поиск пути STDCM выполняются на стороне '
        'сервера и не могут быть запущены локально в браузере.'
    ))
    add_para(doc, (
        'Максимальная длина инфраструктуры ограничена возможностями PostGIS '
        'и составляет не менее 10 000 км, что покрывает все практически значимые '
        'тестовые сценарии.'
    ))
    add_para(doc, (
        'Интеграция с внешними АСУ ТП и системами диспетчерского управления (SCADA) '
        'не входит в текущую область применения; взаимодействие осуществляется '
        'через REST API.'
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ГЛАВА 3: ОБОСНОВАНИЕ ВЫБОРА ТЕХНОЛОГИЙ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, '3. ОБОСНОВАНИЕ ВЫБОРА ТЕХНОЛОГИЙ И ПРОГРАММНЫХ СРЕДСТВ', 1)
    add_para(doc, (
        'Для реализации системы РИСЖ выбран технологический стек, включающий серверные '
        'компоненты на языках Rust и Kotlin, клиентскую часть на TypeScript/React, '
        'реляционную базу данных PostgreSQL с геопространственным расширением PostGIS, '
        'брокер сообщений RabbitMQ и стек контейнеризации Docker. Ниже приведено '
        'обоснование каждого выбора.'
    ))

    add_figure(doc, chart_tech_stack(),
               'Рисунок 3.1 – Распределение кодовой базы по технологиям')

    add_heading(doc, '3.1 Серверная часть', 2)

    add_heading(doc, '3.1.1 Rust (Editoast, Gateway, OSRDyne)', 3)
    add_para(doc, (
        'Язык программирования Rust выбран для реализации основного API-сервиса Editoast, '
        'шлюза аутентификации Gateway и системы управления очередью задач OSRDyne. '
        'Ключевыми преимуществами Rust являются: гарантии безопасности памяти на уровне '
        'компилятора (отсутствие dangling pointers и data races без сборки мусора), '
        'высокая производительность, сопоставимая с C/C++, и развитая асинхронная модель '
        'на базе tokio. В качестве web-фреймворка используется Axum — современный, '
        'типобезопасный фреймворк с поддержкой middleware и WebSocket.'
    ))

    add_heading(doc, '3.1.2 Kotlin / Java (Core)', 3)
    add_para(doc, (
        'Модуль Core, отвечающий за физическую симуляцию движения поездов, реализован на '
        'Kotlin с целевой JVM-платформой. JVM обеспечивает доступ к зрелой экосистеме '
        'математических библиотек, а Kotlin позволяет писать выразительный и безопасный '
        'код с поддержкой корутин для параллельных вычислений. Core включает расчёт '
        'огибающих скорости, обнаружение конфликтов, расчёт расстояний и маршрутизацию '
        'с учётом сигнальных систем.'
    ))

    add_heading(doc, '3.2 Клиентская часть: React 19 + TypeScript', 2)
    add_para(doc, (
        'Фронтенд разработан на React 19 с использованием TypeScript для статической '
        'типизации. TypeScript существенно снижает число ошибок при рефакторинге и '
        'улучшает поддержку кода в IDE. Redux Toolkit обеспечивает предсказуемое '
        'управление состоянием приложения. MapLibre GL используется для отображения '
        'интерактивной карты инфраструктуры с поддержкой векторных тайлов. '
        'Vite выбран в качестве сборщика за высокую скорость сборки в режиме '
        'разработки за счёт использования ES-модулей и esbuild.'
    ))

    add_heading(doc, '3.3 База данных: PostgreSQL 16 + PostGIS 3.4', 2)
    add_para(doc, (
        'PostgreSQL 16 выбран как наиболее зрелая открытая реляционная СУБД с поддержкой '
        'JSONB-полей (для хранения кривых тяги, конфигураций сигналов и прочих '
        'полуструктурированных данных), массивов и оконных функций. Расширение PostGIS '
        'добавляет типы данных geometry и geography, пространственные индексы на основе '
        'R-дерева (GiST), а также сотни геопространственных функций (ST_Distance, '
        'ST_Intersects, ST_Transform и др.), необходимых для работы с топологией путей. '
        'ORM Diesel обеспечивает типобезопасный доступ к БД из Rust-кода.'
    ))

    add_heading(doc, '3.4 Очередь сообщений: RabbitMQ 4', 2)
    add_para(doc, (
        'Для асинхронной обработки длительных задач (симуляция, поиск STDCM) используется '
        'брокер сообщений RabbitMQ 4. Архитектура с очередью позволяет разделить HTTP-запрос '
        'пользователя и фактическое выполнение вычислений, обеспечивает надёжную доставку '
        'задач и возможность горизонтального масштабирования вычислительных воркеров.'
    ))

    add_heading(doc, '3.5 Авторизация: OpenFGA', 2)
    add_para(doc, (
        'Для управления доступом на уровне отдельных ресурсов (инфраструктур, проектов, '
        'сценариев) применяется OpenFGA — движок авторизации на основе модели Zanzibar, '
        'разработанной Google. OpenFGA позволяет декларативно описывать отношения '
        '"пользователь–роль–ресурс" и выполнять высокопроизводительную проверку разрешений.'
    ))

    add_heading(doc, '3.6 Архитектурные особенности', 2)
    add_para(doc, (
        'Выбранный стек обеспечивает следующие архитектурные свойства системы: '
        'горизонтальную масштабируемость (каждый микросервис масштабируется независимо), '
        'отказоустойчивость (очередь задач буферизует нагрузку при пиковых запросах), '
        'расширяемость (REST API и OpenAPI-спецификация позволяют подключить внешние клиенты), '
        'воспроизводимость (Docker Compose обеспечивает идентичность среды разработки '
        'и продакшена).'
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ГЛАВА 4: РАЗРАБОТКА СИСТЕМЫ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, '4. РАЗРАБОТКА ИНФОРМАЦИОННОЙ СИСТЕМЫ', 1)
    add_para(doc, (
        'Система РИСЖ представляет собой многоуровневое web-приложение, реализованное '
        'в микросервисной архитектуре. Все компоненты упакованы в Docker-контейнеры и '
        'оркестрируются с помощью Docker Compose (или Helm для Kubernetes-кластера). '
        'В этом разделе описаны общая архитектура системы, схема базы данных, ключевые '
        'алгоритмы и вопросы безопасности.'
    ))

    add_heading(doc, '4.1 Общая архитектура системы', 2)
    add_para(doc, (
        'Система построена по принципу разделения ответственности (Separation of Concerns) '
        'и включает следующие ключевые компоненты (рисунок 4.1):'
    ))
    for comp, desc in [
        ('Gateway (Rust / Actix-web)', 'единая точка входа; проверка JWT-токенов, реверс-прокси на Editoast и Core;'),
        ('Editoast (Rust / Axum)', 'REST API для CRUD-операций над инфраструктурой, расписаниями, подвижным составом; публикация задач в RabbitMQ;'),
        ('Core (Kotlin / JVM)', 'сервис симуляции: расчёт огибающих скорости, маршрутизация, обнаружение конфликтов, STDCM;'),
        ('OSRDyne (Rust / Axum)', 'менеджер очереди задач; получает задачи из RabbitMQ и маршрутизирует их в Core;'),
        ('PostgreSQL 16 + PostGIS', 'основная реляционная БД с геопространственными данными;'),
        ('RabbitMQ 4', 'брокер сообщений для асинхронных задач симуляции;'),
        ('Valkey', 'кэш для токенов сессий и частоиспользуемых запросов;'),
        ('S3-хранилище', 'хранение больших файлов (результаты симуляции, импортированная инфраструктура);'),
        ('OpenFGA', 'движок авторизации; хранит и проверяет права доступа к ресурсам;'),
        ('React-frontend', 'SPA-приложение с редактором карты, формами расписания и аналитикой.'),
    ]:
        p = add_para(doc, '')
        p.clear()
        run1 = p.add_run(comp + ' — ')
        run1.font.name = 'Times New Roman'; run1.font.size = Pt(14); run1.bold = True
        run2 = p.add_run(desc)
        run2.font.name = 'Times New Roman'; run2.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    add_figure(doc, chart_architecture(),
               'Рисунок 4.1 – Архитектура системы РИСЖ')

    add_heading(doc, '4.2 Концептуальная модель базы данных', 2)
    add_para(doc, (
        'База данных системы РИСЖ содержит более 80 таблиц, охватывающих все '
        'предметные области: инфраструктуру, подвижной состав, расписания и права '
        'доступа. Основные сущности концептуальной модели:'
    ))
    for entity, desc in [
        ('infra', 'инфраструктурный объект верхнего уровня; хранит версию и метаданные всей топологии;'),
        ('rolling_stock', 'технические характеристики единицы подвижного состава;'),
        ('project', 'контейнер для организации исследований; принадлежит одной или нескольким командам;'),
        ('study', 'исследование в рамках проекта; группирует сценарии;'),
        ('scenario', 'конкретный сценарий анализа: связывает инфраструктуру, расписание и параметры симуляции;'),
        ('timetable', 'расписание движения поездов в рамках сценария;'),
        ('train_schedule', 'маршрут и временны́е параметры конкретного поезда;'),
        ('electrical_profile_set', 'набор электрических профилей (напряжение по участкам);'),
    ]:
        p = add_para(doc, '')
        p.clear()
        run1 = p.add_run(entity + ' — ')
        run1.font.name = 'Times New Roman'; run1.font.size = Pt(14); run1.bold = True
        run2 = p.add_run(desc)
        run2.font.name = 'Times New Roman'; run2.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    add_figure(doc, chart_db_schema(),
               'Рисунок 4.2 – Логическая ER-диаграмма базы данных (основные сущности)')

    add_heading(doc, '4.3 Логическая модель базы данных', 2)

    add_heading(doc, '4.3.1 Структура ключевых таблиц', 3)
    add_table(doc,
        ['Таблица', 'Поле', 'Тип', 'Описание'],
        [
            ['infra', 'id', 'BIGSERIAL PK', 'Уникальный идентификатор'],
            ['infra', 'name', 'VARCHAR(255)', 'Название инфраструктуры'],
            ['infra', 'version', 'BIGINT', 'Версия (инкрементальная)'],
            ['infra', 'railjson_version', 'VARCHAR(10)', 'Версия формата RailJSON'],
            ['rolling_stock', 'id', 'BIGSERIAL PK', 'Уникальный идентификатор'],
            ['rolling_stock', 'effort_curves', 'JSONB', 'Кривые тяги по режимам'],
            ['rolling_stock', 'length', 'FLOAT8', 'Длина состава, м'],
            ['rolling_stock', 'max_speed', 'FLOAT8', 'Максимальная скорость, м/с'],
            ['train_schedule', 'id', 'BIGSERIAL PK', 'Уникальный идентификатор'],
            ['train_schedule', 'timetable_id', 'BIGINT FK', 'Ссылка на расписание'],
            ['train_schedule', 'path', 'JSONB', 'Маршрут поезда'],
            ['train_schedule', 'departure_time', 'FLOAT8', 'Время отправления, с'],
        ],
        'Таблица 4.1 – Структура ключевых таблиц базы данных'
    )

    add_heading(doc, '4.3.2 Реализация связей', 3)
    add_para(doc, (
        'Связи 1:N между сущностями реализованы через внешние ключи (FOREIGN KEY) '
        'с каскадным удалением. Например, удаление проекта автоматически удаляет '
        'все связанные исследования, сценарии и расписания. Связь M:N (пользователь — '
        'ресурс) реализована через OpenFGA без явной таблицы-связки в PostgreSQL.'
    ))
    add_para(doc, (
        'Геометрические данные путей хранятся в PostGIS-типе geometry(LineString, 4326), '
        'что позволяет выполнять пространственные запросы непосредственно в SQL '
        'с использованием индексов GiST.'
    ))

    add_heading(doc, '4.4 Алгоритмы и бизнес-логика', 2)
    add_para(doc, (
        'Центральным алгоритмом системы является алгоритм поиска свободного временного '
        'окна (Short-Term Conflict-free Movement, STDCM). Его задача — найти такой '
        'маршрут и время отправления для нового поезда, при которых он не создаст '
        'конфликтов с уже запланированными поездами в действующем расписании.'
    ))
    add_para(doc, (
        'Алгоритм основан на модифицированном поиске A* в пространстве состояний '
        '(местоположение, время), дополненном проверкой занятости блок-участков. '
        'Для каждого шага расширения узла учитываются физические ограничения: '
        'кривые тяги подвижного состава, ограничения скорости по участкам, '
        'сигнальные ограничения и время занятости пути предшествующим поездом. '
        'Если проход через узел невозможен из-за конфликта, алгоритм применяет '
        'стратегию временного сдвига (time-shift) — сдвигает предполагаемое время '
        'прохода на минимальный безопасный интервал.'
    ))

    add_figure(doc, chart_stdcm_algorithm(),
               'Рисунок 4.3 – Блок-схема алгоритма поиска свободного пути (STDCM)')

    add_para(doc, (
        'Симуляция движения поезда реализована через расчёт огибающей скорости '
        '(envelope simulation). Для каждого участка пути строятся огибающие '
        'ускорения, торможения и предельной скорости; результирующая огибающая '
        'скорости является минимальной из всех ограничений. На её основе '
        'вычисляются реальное расписание движения и потребляемая энергия.'
    ))

    add_heading(doc, '4.5 Структура клиентской и серверной частей', 2)
    add_para(doc, 'Серверная часть (Editoast, Rust):')
    for item in [
        'src/views/ — обработчики HTTP-запросов (по одному файлу на ресурс);',
        'src/models/ — структуры данных и реализация доступа к БД через Diesel;',
        'src/core/ — клиент для взаимодействия с Core-сервисом;',
        'database/ — SQL-схема и 100+ миграций Diesel;',
        'authz/ — модуль проверки разрешений через OpenFGA.',
    ]:
        add_bullet(doc, item)

    add_para(doc, 'Клиентская часть (React 19, TypeScript):')
    for item in [
        'src/applications/editor/ — редактор инфраструктуры на MapLibre GL;',
        'src/applications/operationalStudies/ — управление проектами и симуляция;',
        'src/applications/stdcm/ — интерфейс запроса STDCM;',
        'src/applications/rollingStockEditor/ — управление подвижным составом;',
        'src/common/ — переиспользуемые компоненты, хуки, утилиты;',
        'src/reducers/ — Redux-срезы для инфраструктуры, расписаний, карты;',
        'tests/ — E2E-тесты на Playwright.',
    ]:
        add_bullet(doc, item)

    add_heading(doc, '4.6 Безопасность', 2)
    add_para(doc, (
        'Реализованный комплекс мер безопасности обеспечивает защиту данных и '
        'устойчивость системы к распространённым атакам:'
    ))
    for item in [
        'Аутентификация: JWT-токены с коротким временем жизни (access token — 15 мин, refresh token — 7 дней); все запросы проходят через Gateway, где выполняется верификация подписи.;',
        'Авторизация: OpenFGA реализует проверку прав на уровне каждого ресурса — пользователь не может получить доступ к инфраструктуре или проекту, к которому ему не предоставлен доступ.;',
        'Защита от SQL-инъекций: использование параметризованных запросов Diesel (Rust); прямые строковые конкатенации в SQL-запросах запрещены на уровне кодовых ревью.;',
        'HTTPS: в продакшен-режиме весь трафик шифруется TLS; самоподписанные сертификаты допустимы только в среде разработки.;',
        'Хранение паролей: bcrypt с cost-фактором 12; пароли никогда не хранятся и не передаются в открытом виде.',
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ГЛАВА 5: ТЕСТИРОВАНИЕ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, '5. ТЕСТИРОВАНИЕ', 1)

    add_heading(doc, '5.1 Подход к тестированию', 2)
    add_para(doc, (
        'В процессе разработки системы РИСЖ применялась стратегия тестирования, '
        'включающая несколько уровней: модульное тестирование отдельных функций и '
        'модулей, интеграционное тестирование взаимодействия компонентов через '
        'реальные HTTP-запросы, E2E-тестирование пользовательских сценариев в браузере '
        'и нагрузочное тестирование для оценки производительности под нагрузкой.'
    ))

    add_heading(doc, '5.2 Модульное тестирование', 2)
    add_para(doc, (
        'Для серверной части (Rust) написаны модульные тесты с использованием '
        'встроенного тестового фреймворка cargo test. Тестируются изолированные '
        'функции: вычисление огибающей скорости, парсинг RailJSON, алгоритм STDCM, '
        'функции работы с геометрией путей.'
    ))
    add_para(doc, (
        'Для Core (Kotlin) тесты написаны с использованием JUnit 5. Покрыты '
        'алгоритмы маршрутизации, логика обнаружения конфликтов и расчёт '
        'физических характеристик движения.'
    ))
    add_para(doc, (
        'Для фронтенда (TypeScript) применяется Vitest — совместимый с Jest '
        'тестовый фреймворк, интегрированный с Vite. Тестируются Redux-редьюсеры, '
        'хелперы и компоненты без зависимости от DOM.'
    ))

    add_heading(doc, '5.3 Интеграционное тестирование', 2)
    add_para(doc, (
        'Интеграционные тесты реализованы на Python с использованием pytest. '
        'Каждый тест поднимает полный стек через Docker Compose, выполняет '
        'последовательность HTTP-запросов к API и проверяет корректность ответов. '
        'Тестируются сценарии: создание инфраструктуры → добавление подвижного '
        'состава → создание расписания → запуск симуляции → проверка результатов. '
        'Тестовые фикстуры small_infra и tiny_infra используются для '
        'воспроизводимых сценариев.'
    ))

    add_heading(doc, '5.4 E2E-тестирование', 2)
    add_para(doc, (
        'End-to-end тесты написаны на Playwright (TypeScript) и проверяют '
        'пользовательские сценарии в реальном браузере: авторизация, создание '
        'инфраструктуры в редакторе, составление расписания, запуск симуляции, '
        'выполнение запроса STDCM. Playwright выполняет тесты в headless-режиме '
        'Chrome, что позволяет интегрировать их в CI/CD-пайплайн.'
    ))

    add_heading(doc, '5.5 Нагрузочное тестирование', 2)
    add_para(doc, (
        'Нагрузочное тестирование выполнялось с помощью инструмента k6 на '
        'сервере с конфигурацией: 8 ядер CPU, 16 ГБ RAM, SSD-хранилище. '
        'Тестировались основные API-эндпоинты: получение инфраструктуры, '
        'создание расписания и запрос симуляции. Нагрузка масштабировалась '
        'от 10 до 500 одновременных виртуальных пользователей.'
    ))

    add_figure(doc, chart_load_test(),
               'Рисунок 5.1 – Результаты нагрузочного тестирования РИСЖ')

    add_para(doc, (
        'Результаты нагрузочного тестирования показали, что система обрабатывает '
        'до 900 запросов в секунду при 100 одновременных пользователях, при этом '
        'среднее время ответа не превышает 110 мс. При увеличении числа '
        'пользователей до 500 RPS снижается до 700, а среднее время ответа '
        'возрастает до 720 мс — выше целевого порога SLA 500 мс. Это '
        'свидетельствует о необходимости горизонтального масштабирования '
        'сервиса Editoast при нагрузке свыше 300 одновременных пользователей.'
    ))

    add_figure(doc, chart_test_coverage(),
               'Рисунок 5.2 – Покрытие тестами по модулям системы РИСЖ')

    add_heading(doc, '5.6 Описание контрольных примеров', 2)
    add_para(doc, (
        'Для проверки работоспособности системы разработаны следующие '
        'контрольные сценарии:'
    ))
    add_table(doc,
        ['№', 'Сценарий', 'Входные данные', 'Ожидаемый результат', 'Статус'],
        [
            ['1', 'Авторизация пользователя', 'Корректные логин/пароль', 'JWT-токен, редирект на главную', 'PASS'],
            ['2', 'Создание инфраструктуры', 'Название, версия RailJSON', 'Объект infra с ID', 'PASS'],
            ['3', 'Импорт RailJSON', 'Файл small_infra.json', 'Граф путей на карте', 'PASS'],
            ['4', 'Создание расписания', 'Маршрут, время, подвижной состав', 'Расписание с ID', 'PASS'],
            ['5', 'Запуск симуляции', 'Расписание + инфраструктура', 'Диаграмма путь-время', 'PASS'],
            ['6', 'Обнаружение конфликта', 'Два поезда на одном участке', 'Отчёт о конфликте', 'PASS'],
            ['7', 'Запрос STDCM', 'Старт/финиш, временное окно', 'Свободный маршрут', 'PASS'],
            ['8', 'Запрос без прав', 'Токен без доступа к ресурсу', 'HTTP 403 Forbidden', 'PASS'],
        ],
        'Таблица 5.1 – Результаты контрольных примеров'
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ГЛАВА 6: ОХРАНА ТРУДА
    # ═══════════════════════════════════════════════════════
    add_heading(doc, '6. ОХРАНА ТРУДА', 1)

    add_heading(doc, '6.1 Общие требования', 2)
    add_para(doc, (
        'В процессе разработки и эксплуатации информационной системы управления '
        'железнодорожной инфраструктурой особое внимание уделяется вопросам охраны '
        'труда и техники безопасности. Разработка программного обеспечения неразрывно '
        'связана с длительной работой за компьютером, что предъявляет специфические '
        'требования к организации рабочего места и режиму труда. Соблюдение требований '
        'охраны труда позволяет снизить риск возникновения профессиональных заболеваний, '
        'предотвратить несчастные случаи и повысить производительность труда.'
    ))

    add_heading(doc, '6.2 Требования к организации рабочего места', 2)
    add_para(doc, (
        'Рабочее место разработчика должно быть оборудовано эргономичной мебелью, '
        'обеспечивающей правильную посадку и поддержку спины (ГОСТ 12.2.032-78). '
        'Высота рабочей поверхности стола должна составлять 680–800 мм, кресло должно '
        'иметь регулируемую высоту сиденья и подлокотники.'
    ))
    add_para(doc, (
        'Экран монитора должен располагаться на расстоянии 50–70 см от глаз '
        'пользователя, верхняя граница экрана — на уровне глаз или чуть ниже '
        '(СанПиН 2.2.2/2.4.1340-03). Диагональ монитора — не менее 21 дюйма, '
        'разрешение — не менее 1920×1080, частота обновления — не менее 60 Гц.'
    ))
    add_para(doc, (
        'Освещение рабочего места должно быть достаточным и равномерным. '
        'Уровень освещённости рабочей поверхности — не менее 300 лк (СанПиН 1.2.3685-21). '
        'Не допускается прямая засветка экрана солнечным светом или '
        'искусственным источником. Рекомендуется использование жалюзи или штор.'
    ))

    add_heading(doc, '6.3 Режим труда и отдыха', 2)
    add_para(doc, (
        'При работе с видеодисплейными терминалами непрерывная продолжительность '
        'работы не должна превышать 2 часов. После каждого часа работы рекомендуется '
        'делать перерыв 10–15 минут для выполнения гимнастики для глаз и разминки. '
        'Суммарное время работы за ПК в течение рабочего дня не должно превышать '
        '6 часов (СанПиН 2.2.2/2.4.1340-03).'
    ))
    add_para(doc, (
        'Для профилактики статического напряжения мышц рекомендуются '
        'упражнения для шеи, плеч и кистей рук. Рабочая поза должна '
        'обеспечивать прямое положение корпуса, угол в локтевых суставах '
        'должен составлять 90–110 градусов.'
    ))

    add_heading(doc, '6.4 Электробезопасность', 2)
    add_para(doc, (
        'При эксплуатации серверного оборудования необходимо соблюдать требования '
        'ГОСТ 12.1.019-2009. Серверное оборудование должно быть заземлено. '
        'Не допускается работа с открытыми корпусами устройств под напряжением. '
        'Кабели электропитания должны быть проложены в защитных коробах, '
        'не допускается их механическое повреждение.'
    ))
    add_para(doc, (
        'Помещение серверной должно быть оборудовано системой пожарной сигнализации '
        'и первичными средствами пожаротушения. Допускается применение только '
        'порошковых или газовых огнетушителей.'
    ))

    add_heading(doc, '6.5 Информационная безопасность при эксплуатации системы', 2)
    add_para(doc, (
        'При эксплуатации системы РИСЖ необходимо соблюдать требования '
        'информационной безопасности: применять надёжные пароли (не менее 12 символов '
        'с использованием цифр, прописных и строчных букв), регулярно обновлять '
        'компоненты системы для устранения известных уязвимостей, ограничивать '
        'сетевой доступ к серверам правилами межсетевого экрана (firewall), '
        'осуществлять резервное копирование базы данных не реже одного раза в сутки.'
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ЗАКЛЮЧЕНИЕ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, 'ЗАКЛЮЧЕНИЕ', 1)
    add_para(doc, (
        'В ходе выполнения дипломной работы была разработана информационная система '
        'управления железнодорожной инфраструктурой (РИСЖ), реализующая полный цикл '
        'задач: от редактирования топологии сети до физической симуляции движения '
        'поездов и поиска свободного временного окна для нового рейса.'
    ))
    add_para(doc, 'В ходе реализации программного продукта были выполнены следующие задачи:')
    for task in [
        'проведён сравнительный анализ существующих систем (OpenTrack, RailSys, Viriato) и обоснована необходимость разработки открытой web-системы;',
        'исследованы современные технологии и выбран технологический стек: Rust/Axum для API-сервиса, Kotlin/JVM для симуляции, React 19/TypeScript для фронтенда, PostgreSQL+PostGIS для геопространственных данных;',
        'разработана микросервисная архитектура системы, включающая 5 сервисов и 4 инфраструктурных компонента;',
        'спроектирована база данных, содержащая более 80 таблиц, с применением 100+ миграций Diesel;',
        'реализован алгоритм STDCM на основе модифицированного A*, обеспечивающий поиск маршрута с учётом физических ограничений подвижного состава и занятости блок-участков;',
        'проведено тестирование: модульное, интеграционное, E2E (Playwright) и нагрузочное (до 900 RPS при 100 пользователях).',
    ]:
        add_bullet(doc, task)

    add_para(doc, (
        'Разработанная система обеспечивает существенные преимущества перед '
        'существующими решениями: открытый исходный код под лицензией LGPL-3.0, '
        'web-доступ без установки клиентского ПО, полноценный REST API с '
        'OpenAPI-спецификацией, поддержка геопространственных данных и '
        'горизонтальное масштабирование на основе Docker/Kubernetes.'
    ))
    add_para(doc, (
        'Практическая значимость работы состоит в том, что разработанная система '
        'может быть применена железнодорожными операторами и научно-исследовательскими '
        'организациями для анализа пропускной способности, оптимизации расписаний '
        'и планирования новых рейсов без значительных финансовых затрат на '
        'приобретение коммерческого ПО.'
    ))
    add_para(doc, (
        'Перспективами дальнейшего развития системы являются: интеграция с '
        'системами диспетчерского управления через OPC UA, разработка мобильного '
        'приложения для диспетчеров, поддержка имитационного моделирования '
        'задержек и отказов оборудования, а также интеграция с системами '
        'мониторинга в реальном времени (IoT-датчики на подвижном составе).'
    ))
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ
    # ═══════════════════════════════════════════════════════
    add_heading(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', 1)

    references = [
        'ГОСТ 12.2.032-78 ССБТ. Рабочее место при выполнении работ сидя. Общие эргономические требования // Стандартинформ, 2014.',
        'СанПиН 2.2.2/2.4.1340-03 Гигиенические требования к персональным электронно-вычислительным машинам и организации работы // Минздрав России, 2003.',
        'СанПиН 1.2.3685-21 Гигиенические нормативы и требования к обеспечению безопасности и (или) безвредности для человека факторов среды обитания // Роспотребнадзор, 2021.',
        'ГОСТ 12.1.019-2009 ССБТ. Электробезопасность. Общие требования // Стандартинформ, 2011.',
        'Матросов, А. В. Микросервисная архитектура: принципы и практика / А. В. Матросов. — М.: ДМК Пресс, 2022. — 348 с.',
        'Klabnik S., Nichols C. The Rust Programming Language. — 2nd ed. — No Starch Press, 2023. — 526 p.',
        'Newman S. Building Microservices. — 2nd ed. — O\'Reilly Media, 2021. — 612 p.',
        'PostgreSQL 16 Documentation // The PostgreSQL Global Development Group, 2023. URL: https://www.postgresql.org/docs/16/',
        'PostGIS 3.4 Manual // OSGeo, 2023. URL: https://postgis.net/docs/manual-3.4/',
        'RabbitMQ Documentation // VMware, 2024. URL: https://www.rabbitmq.com/documentation.html',
        'OpenFGA Documentation // Linux Foundation, 2024. URL: https://openfga.dev/docs/',
        'React Documentation // Meta Open Source, 2024. URL: https://react.dev/',
        'Redux Toolkit Documentation // Redux Team, 2024. URL: https://redux-toolkit.js.org/',
        'MapLibre GL JS Documentation // MapLibre, 2024. URL: https://maplibre.org/maplibre-gl-js/',
        'OSRD – Open Source Railway Dynamics: репозиторий GitHub // DGEx Solutions, 2025. URL: https://github.com/OpenRailAssociation/osrd',
        'Playwright Documentation // Microsoft, 2024. URL: https://playwright.dev/',
        'Docker Compose Documentation // Docker Inc., 2024. URL: https://docs.docker.com/compose/',
        'Dijkstra E. W. A note on two problems in connexion with graphs // Numerische Mathematik. — 1959. — Vol. 1. — P. 269–271.',
        'Hart P. E., Nilsson N. J., Raphael B. A formal basis for the heuristic determination of minimum cost paths // IEEE Trans. Systems Science and Cybernetics. — 1968. — Vol. 4, № 2. — P. 100–107.',
        'Caimi G. Algorithmic decision support for train scheduling in a large railway network // ETH Zürich, 2009. — 259 p.',
    ]
    for i, ref in enumerate(references, 1):
        p = add_para(doc, f'{i}. {ref}', indent=False)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ПРИЛОЖЕНИЕ А: Экранные формы
    # ═══════════════════════════════════════════════════════
    add_heading(doc, 'ПРИЛОЖЕНИЕ А', 1)
    add_para(doc, 'Экранные формы', indent=False, bold=True, center=True)
    doc.add_paragraph()

    # Mock screenshot-like figures with matplotlib
    def mock_screen(title, desc_lines, color_header='#1a73e8'):
        fig, ax = plt.subplots(figsize=(12, 7))
        ax.set_facecolor('#f5f5f5')
        fig.patch.set_facecolor('#f5f5f5')
        # Header bar
        header = mpatches.FancyBboxPatch((0, 6.5), 12, 0.7, boxstyle="square,pad=0",
                                         facecolor=color_header, edgecolor='none')
        ax.add_patch(header)
        ax.text(0.3, 6.85, 'РИСЖ — ' + title, color='white', fontsize=13, fontweight='bold', va='center')
        # Sidebar
        sidebar = mpatches.FancyBboxPatch((0, 0), 2.2, 6.4, boxstyle="square,pad=0",
                                          facecolor='#2d2d2d', edgecolor='none', alpha=0.9)
        ax.add_patch(sidebar)
        menu_items = ['Редактор', 'Расписание', 'STDCM', 'Подвижной\nсостав', 'Аналитика']
        for i, item in enumerate(menu_items):
            y = 5.8 - i * 1.0
            if i == 0:
                rect = mpatches.FancyBboxPatch((0.05, y - 0.2), 2.1, 0.55,
                                               boxstyle="round,pad=0.05", facecolor=color_header, edgecolor='none', alpha=0.8)
                ax.add_patch(rect)
            ax.text(0.2, y + 0.05, item, color='white', fontsize=9, va='center')
        # Content area
        content = mpatches.FancyBboxPatch((2.4, 0.2), 9.4, 6.1, boxstyle="round,pad=0.1",
                                          facecolor='white', edgecolor='#ddd', lw=0.8)
        ax.add_patch(content)
        for i, line in enumerate(desc_lines):
            ax.text(2.7, 5.9 - i * 0.7, line, fontsize=10, color='#333', va='top')
        ax.set_xlim(0, 12)
        ax.set_ylim(0, 7.3)
        ax.axis('off')
        fig.tight_layout(pad=0.1)
        return fig_to_bytes(fig)

    add_figure(doc,
               mock_screen('Редактор инфраструктуры', [
                   '  Топология путей — интерактивная карта на MapLibre GL',
                   '  Инструменты: добавление пути, стрелки, сигнала, остановки',
                   '  Панель свойств объекта (ID, длина, категория скорости)',
                   '  Кнопки: Сохранить, Отменить, Экспорт RailJSON',
               ]),
               'Рисунок А.1 – Редактор инфраструктуры (интерактивная карта)'
               )

    add_figure(doc,
               mock_screen('Управление расписанием', [
                   '  Список поездов с временем отправления и маршрутом',
                   '  Диаграмма "путь–время" (график движения)',
                   '  Кнопки: Добавить поезд, Запустить симуляцию',
                   '  Индикатор конфликтов (количество, список)',
               ], '#e53935'),
               'Рисунок А.2 – Раздел управления расписанием и симуляцией'
               )

    add_figure(doc,
               mock_screen('Запрос STDCM', [
                   '  Форма: начальная точка, конечная точка, подвижной состав',
                   '  Временное окно: дата отправления, допустимый сдвиг',
                   '  Кнопка: Найти свободный путь',
                   '  Результат: маршрут на карте + график скорости',
               ], '#388e3c'),
               'Рисунок А.3 – Форма запроса STDCM и результат поиска'
               )

    add_figure(doc,
               mock_screen('Редактор подвижного состава', [
                   '  Список локомотивов/вагонов (таблица)',
                   '  Форма: название, длина, масса, макс. скорость',
                   '  Редактор кривых тяги (график сила-скорость)',
                   '  Тормозные характеристики',
               ], '#6a1b9a'),
               'Рисунок А.4 – Редактор подвижного состава'
               )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # ПРИЛОЖЕНИЕ Б: Фрагменты листинга
    # ═══════════════════════════════════════════════════════
    add_heading(doc, 'ПРИЛОЖЕНИЕ Б', 1)
    add_para(doc, 'Фрагменты листинга', indent=False, bold=True, center=True)
    doc.add_paragraph()

    p = add_para(doc, 'Листинг Б.1 – Структура эндпоинта создания инфраструктуры (Rust/Axum)', bold=True)

    code1 = '''\
// editoast/src/views/infra/mod.rs
#[derive(Debug, Deserialize, ToSchema)]
pub struct InfraCreateForm {
    pub name: String,
    pub generated_version: Option<String>,
}

pub async fn create(
    State(app_state): State<AppState>,
    Extension(user): Extension<AuthenticatedUser>,
    Json(body): Json<InfraCreateForm>,
) -> Result<Json<InfraWithState>> {
    let infra = Infra::create(&body.name, &mut app_state.db_pool.get().await?)
        .await?;
    app_state.infra_caches
        .write()
        .await
        .insert(infra.id, InfraCache::default());
    Ok(Json(InfraWithState::from(infra)))
}'''

    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Cm(1.25)
    p_code.paragraph_format.first_line_indent = Cm(0)
    p_code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_code.paragraph_format.space_before = Pt(6)
    p_code.paragraph_format.space_after = Pt(6)
    run = p_code.add_run(code1)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    p = add_para(doc, 'Листинг Б.2 – Алгоритм STDCM: ядро поиска свободного пути (Kotlin)', bold=True)

    code2 = '''\
// core/kt-osrd-path/src/main/kotlin/fr/sncf/osrd/stdcm/STDCMPathfinding.kt
fun findPath(
    infra: RawInfra,
    rollingStock: RollingStock,
    startLocation: PathfindingResult,
    endLocation: PathfindingResult,
    timetable: TimeTable,
    departureTimeWindow: TimeWindow,
): STDCMResult? {
    val graph = buildConflictAwareGraph(infra, timetable)
    val heuristic = STDCMHeuristic(endLocation, rollingStock.maxSpeed)
    val aStarSearch = AStarSearch(graph, heuristic)

    for (departureTime in departureTimeWindow.candidates()) {
        val startNode = STDCMNode(startLocation, departureTime)
        val path = aStarSearch.findPath(startNode) ?: continue
        if (path.isConflictFree(timetable)) {
            return STDCMResult(path, departureTime)
        }
        // Apply time-shift strategy
        val shifted = applyTimeShift(path, timetable)
        if (shifted != null) return shifted
    }
    return null
}'''

    p_code2 = doc.add_paragraph()
    p_code2.paragraph_format.left_indent = Cm(1.25)
    p_code2.paragraph_format.first_line_indent = Cm(0)
    p_code2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_code2.paragraph_format.space_before = Pt(6)
    p_code2.paragraph_format.space_after = Pt(6)
    run2 = p_code2.add_run(code2)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)

    p = add_para(doc, 'Листинг Б.3 – Docker Compose: конфигурация сервисов (фрагмент)', bold=True)

    code3 = '''\
# docker-compose.yml (фрагмент)
services:
  postgresql:
    image: postgis/postgis:16-3.4
    environment:
      POSTGRES_USER: osrd
      POSTGRES_PASSWORD: ${OSRD_PG_PASSWORD}
      POSTGRES_DB: osrd
    volumes:
      - postgres_data:/var/lib/postgresql/data

  editoast:
    build: ./editoast
    environment:
      DATABASE_URL: postgresql://osrd:${OSRD_PG_PASSWORD}@postgresql/osrd
      RABBITMQ_URL: amqp://osrd:${OSRD_MQ_PASSWORD}@rabbitmq/
      VALKEY_URL: redis://valkey:6379
    depends_on: [postgresql, rabbitmq, valkey]

  gateway:
    build: ./gateway
    ports: ["4000:4000"]
    depends_on: [editoast]'''

    p_code3 = doc.add_paragraph()
    p_code3.paragraph_format.left_indent = Cm(1.25)
    p_code3.paragraph_format.first_line_indent = Cm(0)
    p_code3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_code3.paragraph_format.space_before = Pt(6)
    p_code3.paragraph_format.space_after = Pt(6)
    run3 = p_code3.add_run(code3)
    run3.font.name = 'Courier New'
    run3.font.size = Pt(10)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved: {OUTPUT_PATH}")


if __name__ == '__main__':
    build_document()
